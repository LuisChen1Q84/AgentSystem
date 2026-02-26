#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Markdown to DOCX Converter for 513180 Research Report
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re
import os

# File paths
INPUT_FILE = "/Volumes/Luis_MacData/AgentSystem/docs/513180_report_verified.md"
OUTPUT_FILE = "/Volumes/Luis_MacData/AgentSystem/产出/513180_恒生科技ETF_研究报告.docx"

# Chart paths
CHART_PATHS = {
    "performance": "/Volumes/Luis_MacData/AgentSystem/charts/01_performance.png",
    "sector_distribution": "/Volumes/Luis_MacData/AgentSystem/charts/02_sector_distribution.png",
    "top_holdings": "/Volumes/Luis_MacData/AgentSystem/charts/03_top_holdings.png",
    "fund_flow": "/Volumes/Luis_MacData/AgentSystem/charts/04_fund_flow.png",
}

def set_cell_shading(cell, color="FFFFFF"):
    """Set cell background color"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading_elm)

def add_formatted_heading(doc, text, level=1):
    """Add a formatted heading"""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = 'SimSun'
        run.font.size = Pt(16 if level == 1 else 14)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 51, 102)
    return heading

def add_formatted_paragraph(doc, text, bold=False, font_size=11):
    """Add a formatted paragraph"""
    para = doc.add_paragraph(text)
    for run in para.runs:
        run.font.name = 'SimSun'
        run.font.size = Pt(font_size)
        run.font.bold = bold
    return para

def read_markdown(file_path):
    """Read markdown file"""
    with open(file_path, 'r', encoding='utf-8') as f:
        return f.read()

def parse_markdown_content(content):
    """Parse markdown content into structured data"""
    lines = content.split('\n')
    sections = []
    current_section = None
    current_subsection = None
    current_content = []

    for line in lines:
        # Skip horizontal rules
        if line.strip() == '---':
            continue

        # Main heading (##)
        if line.startswith('## '):
            if current_section:
                sections.append(current_section)
            current_section = {
                'title': line.replace('## ', '').strip(),
                'level': 1,
                'subsections': [],
                'content': []
            }
            current_content = current_section['content']
            current_subsection = None
        # Subsection (###)
        elif line.startswith('### '):
            if current_subsection:
                current_section['subsections'].append(current_subsection)
            current_subsection = {
                'title': line.replace('### ', '').strip(),
                'level': 2,
                'content': []
            }
            current_content = current_subsection['content']
        # Content line
        elif line.strip():
            current_content.append(line.strip())
        # Empty line
        elif current_content and current_content[-1] != '':
            current_content.append('')

    if current_section:
        if current_subsection:
            current_section['subsections'].append(current_subsection)
        sections.append(current_section)

    return sections

def add_table(doc, table_data, style='Table Grid'):
    """Add a table to the document"""
    if not table_data:
        return None

    rows = len(table_data)
    cols = len(table_data[0])

    table = doc.add_table(rows=rows, cols=cols)
    table.style = style
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, row_data in enumerate(table_data):
        row = table.rows[i]
        for j, cell_data in enumerate(row_data):
            cell = row.cells[j]
            cell.text = str(cell_data)
            # Format header row
            if i == 0:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = 'SimSun'
                        run.font.size = Pt(10)
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(255, 255, 255)
                set_cell_shading(cell, "003366")
            else:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = 'SimSun'
                        run.font.size = Pt(9)

    return table

def process_table_content(lines, start_idx):
    """Process table content from markdown"""
    table_data = []
    i = start_idx

    while i < len(lines):
        line = lines[i].strip()
        if line.startswith('|') and line.endswith('|'):
            # Parse table row
            cells = [cell.strip() for cell in line.split('|')[1:-1]]
            # Skip separator rows (containing ---)
            if not any('---' in cell for cell in cells):
                table_data.append(cells)
            i += 1
        else:
            break

    return table_data, i

def create_document():
    """Create the DOCX document"""
    # Read markdown
    content = read_markdown(INPUT_FILE)
    sections = parse_markdown_content(content)

    # Create document
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'SimSun'
    style.font.size = Pt(11)

    # Title
    title = doc.add_heading('513180 恒生科技指数ETF 深度研究报告', level=0)
    for run in title.runs:
        run.font.name = 'SimSun'
        run.font.size = Pt(22)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 51, 102)

    # Date
    date_para = doc.add_paragraph()
    date_para.add_run('研究报告日期：2026年2月26日').italic = True

    # Process sections
    section_idx = 0
    for section in sections:
        # Skip if it's not a main section we want
        if section['title'] in ['执行摘要', '一、产品概述', '二、业绩分析', '三、持仓分析',
                                 '四、估值分析', '五、资金流向分析', '六、风险分析',
                                 '七、投资建议', '八、结论', '来源']:
            section_idx += 1

        # Add main section heading
        if section['level'] == 1:
            add_formatted_heading(doc, section['title'], level=1)

        # Process content before subsections
        for line in section['content']:
            if line.startswith('|'):
                # Table
                table_data, _ = process_table_content([line], 0)
                # Try to get more rows
                content_idx = section['content'].index(line)
                for i in range(content_idx + 1, len(section['content'])):
                    if section['content'][i].startswith('|'):
                        more_data, _ = process_table_content([section['content'][i]], 0)
                        table_data.extend(more_data)
                    else:
                        break

                if table_data:
                    add_table(doc, table_data)
            elif line.startswith('**') and line.endswith('**'):
                # Bold key point
                add_formatted_paragraph(doc, line.replace('**', ''), bold=True)
            elif line and not line.startswith('#'):
                add_formatted_paragraph(doc, line)

        # Process subsections
        for subsection in section['subsections']:
            # Add subsection heading
            add_formatted_heading(doc, subsection['title'], level=2)

            # Add chart based on section/subsection title
            chart_to_add = None

            # Section 2: 业绩分析 -> performance chart
            if '二、业绩分析' in section['title'] or section_idx == 2:
                if '2.1' in subsection['title'] or '历史业绩' in subsection['title']:
                    chart_to_add = 'performance'

            # Section 3: 持仓分析 -> sector and holdings charts
            if '三、持仓分析' in section['title'] or section_idx == 3:
                if '行业分布' in subsection['title']:
                    chart_to_add = 'sector_distribution'
                elif '重仓' in subsection['title']:
                    chart_to_add = 'top_holdings'

            # Section 5: 资金流向 -> fund flow chart
            if '五、资金流向' in section['title'] or section_idx == 5:
                if '5.1' in subsection['title'] or '南向' in subsection['title']:
                    chart_to_add = 'fund_flow'

            # Add chart if identified
            if chart_to_add and CHART_PATHS.get(chart_to_add):
                try:
                    if os.path.exists(CHART_PATHS[chart_to_add]):
                        doc.add_picture(CHART_PATHS[chart_to_add], width=Inches(5))
                        # Add caption
                        caption = doc.add_paragraph()
                        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        caption_run = caption.add_run(f"图表: {chart_to_add.replace('_', ' ').title()}")
                        caption_run.font.name = 'SimSun'
                        caption_run.font.size = Pt(9)
                        caption_run.font.italic = True
                except Exception as e:
                    print(f"Warning: Could not add chart {chart_to_add}: {e}")

            # Process subsection content
            for line in subsection['content']:
                if line.startswith('|'):
                    # Table
                    table_data = []
                    content_idx = subsection['content'].index(line)
                    for i in range(content_idx, len(subsection['content'])):
                        row_line = subsection['content'][i].strip()
                        if row_line.startswith('|') and row_line.endswith('|'):
                            cells = [cell.strip() for cell in row_line.split('|')[1:-1]]
                            if not any('---' in cell for cell in cells):
                                table_data.append(cells)
                        else:
                            break

                    if table_data:
                        add_table(doc, table_data)
                elif line.startswith('**') and line.endswith('**'):
                    add_formatted_paragraph(doc, line.replace('**', ''), bold=True)
                elif line and not line.startswith('#'):
                    add_formatted_paragraph(doc, line)

    # Add page breaks between major sections
    # This is done by adding section breaks

    # Save document
    doc.save(OUTPUT_FILE)
    print(f"Document saved to: {OUTPUT_FILE}")

if __name__ == "__main__":
    # Ensure output directory exists
    os.makedirs(os.path.dirname(OUTPUT_FILE), exist_ok=True)
    create_document()
